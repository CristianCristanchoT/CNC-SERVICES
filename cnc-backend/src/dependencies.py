import os, shutil
import glob
import json
from bson import json_util
import string
import random
from docx import Document
from docx.shared import Inches
from fastapi.responses import FileResponse
from datetime import datetime

from .settings import (
    preguntas
)

def parse_json(data):
    return json.loads(json_util.dumps(data))

def get_sectores():
    sectores = list(preguntas.distinct("Sector"))
    return sectores

def get_segmentos():
    segmentos = list(preguntas.distinct("Segmento"))
    return segmentos

def get_clientes():
    clientes = list(preguntas.distinct("Cliente"))
    return clientes

def incluir_cliente_preguntas(preguntas, cliente):

    for pregunta in preguntas:

        if 'Cliente' in pregunta['Preguntas']:
                
                pregunta['Preguntas'] = pregunta['Preguntas'].replace('Cliente', cliente)

    return preguntas

def incluir_competidor_preguntas(preguntas, competidor):

    for pregunta in preguntas:

        if 'Competidor' in pregunta['Preguntas']:

             pregunta['Preguntas'] = pregunta['Preguntas'].replace('Competidor', competidor)

    return preguntas

def excluir_preguntas_competidor(preguntas):

    result = []

    for pregunta in preguntas:

        if 'Competidor' not in pregunta['Preguntas']:
                
            result.append(pregunta)  

    return result


def obtener_preguntas(sector  = '', segmento = '', cliente = '', competidor = '', solo_historicos = True, sin_competidor = False):
    """
    print('-'*10)
    print('sector: ' + sector)
    print('segmento: ' + segmento)
    print('cliente: ' + cliente)
    print('competidor: ' + competidor)
    print('solo_historicos: ' + str(solo_historicos))
    print('sin_competidor: ' + str(sin_competidor))
    print('-'*10)
    """
    ###############################################################################################
    if sector and segmento and solo_historicos:

        #print('sector and segmento and solo_historico')

        preguntas_busqueda = list(preguntas.aggregate([
            { "$match":{
                    "$and": [
                        {"Sector": { "$eq" : sector}},
                        {"Segmento" : { "$eq" : segmento}},
                        {"Cliente" : { "$eq" : cliente}}
                    ]
                }
            },
            { "$sort": { "Grupos" : 1 }}
        ]))

        preguntas_busqueda = parse_json(preguntas_busqueda)

        preguntas_busqueda = incluir_cliente_preguntas(preguntas_busqueda, cliente)
        
        if sin_competidor:
            preguntas_busqueda = excluir_preguntas_competidor(preguntas_busqueda)
        else: 
            preguntas_busqueda = incluir_competidor_preguntas(preguntas_busqueda, competidor)

        return preguntas_busqueda

    ###############################################################################################
    if sector and solo_historicos:

        #print('sector and solo_historico')

        preguntas_busqueda = list(preguntas.aggregate([
            { "$match":{
                    "$and": [
                        {"Sector": { "$eq" : sector}},
                        {"Cliente" : { "$eq" : cliente}}
                    ]
                }
            },
            { "$sort": { "Grupos" : 1 }}
        ]))

        preguntas_busqueda = parse_json(preguntas_busqueda)

        preguntas_busqueda = incluir_cliente_preguntas(preguntas_busqueda, cliente)

        if sin_competidor:
            preguntas_busqueda = excluir_preguntas_competidor(preguntas_busqueda)
        else: 
            preguntas_busqueda = incluir_competidor_preguntas(preguntas_busqueda, competidor)

        return preguntas_busqueda

    ###############################################################################################
    if segmento and solo_historicos:

        #print('segmento and solo_historico')

        preguntas_busqueda = list(preguntas.aggregate([
            { "$match":{
                    "$and": [
                        {"Segmento" : { "$eq" : segmento}},
                        {"Cliente" : { "$eq" : cliente}}
                    ]
                }
            },
            { "$sort": { "Grupos" : 1 }}
        ]))

        preguntas_busqueda = parse_json(preguntas_busqueda)

        preguntas_busqueda = incluir_cliente_preguntas(preguntas_busqueda, cliente)

        if sin_competidor:
            preguntas_busqueda = excluir_preguntas_competidor(preguntas_busqueda)
        else: 
            preguntas_busqueda = incluir_competidor_preguntas(preguntas_busqueda, competidor)

        return preguntas_busqueda

    ###############################################################################################
    if sector and segmento:

        #print('sector and segmento')

        preguntas_busqueda = list(preguntas.aggregate([
            { "$match":{
                    "$and": [
                        {"Sector": { "$eq" : sector}},
                        {"Segmento" : { "$eq" : segmento}}
                    ]
                }
            },
            { "$sort": { "Grupos" : 1 }}
        ]))

        preguntas_busqueda = parse_json(preguntas_busqueda)

        preguntas_busqueda = incluir_cliente_preguntas(preguntas_busqueda, cliente)

        if sin_competidor:
            preguntas_busqueda = excluir_preguntas_competidor(preguntas_busqueda)
        else: 
            preguntas_busqueda = incluir_competidor_preguntas(preguntas_busqueda, competidor)

        return preguntas_busqueda

    ###############################################################################################
    if sector:

        #print('sector')

        preguntas_busqueda = list(preguntas.aggregate([
            { "$match":{
                    "$and": [
                        {"Sector": { "$eq" : sector}}
                    ]
                }
            },
            { "$sort": { "Grupos" : 1 }}
        ]))

        preguntas_busqueda = parse_json(preguntas_busqueda)

        preguntas_busqueda = incluir_cliente_preguntas(preguntas_busqueda, cliente)

        if sin_competidor:
            preguntas_busqueda = excluir_preguntas_competidor(preguntas_busqueda)
        else: 
            preguntas_busqueda = incluir_competidor_preguntas(preguntas_busqueda, competidor)

        return preguntas_busqueda

    ###############################################################################################
    if segmento:

        #print('segmento')

        preguntas_busqueda = list(preguntas.aggregate([
            { "$match":{
                    "$and": [
                        {"Segmento" : { "$eq" : segmento}}
                    ]
                }
            },
            { "$sort": { "Grupos" : 1 }}
        ]))

        preguntas_busqueda = parse_json(preguntas_busqueda)

        preguntas_busqueda = incluir_cliente_preguntas(preguntas_busqueda, cliente)

        if sin_competidor:
            preguntas_busqueda = excluir_preguntas_competidor(preguntas_busqueda)
        else: 
            preguntas_busqueda = incluir_competidor_preguntas(preguntas_busqueda, competidor)

        return preguntas_busqueda


def generar_documento_word(metadatos, preguntas_seleccionadas):

    document_name = id_generator()

    secciones_encuesta = obtener_inidices_orden_de_las_preguntas(preguntas_seleccionadas)

    document = Document()

    index_pregunta = 1

    for i in range(len(secciones_encuesta)): 

        titulo_seccion = int_to_Roman(i+1) + '. ' +  secciones_encuesta[i].capitalize()

        document.add_paragraph(' ')
        document = generar_titulo(document, titulo_seccion)

        for pregunta in preguntas_seleccionadas:

            if pregunta['Categorias'] == secciones_encuesta[i]:

                texto_pregunta = str(index_pregunta) + '. ' + pregunta['Preguntas']
                index_pregunta += 1

                document.add_paragraph(texto_pregunta)

                if pregunta['Escala'] == 'SI':
                    
                    try:

                        tabla_image_path = 'docs/images/' + pregunta['Tipo de escala'] + '.JPG'
                        document.add_picture(tabla_image_path, width=Inches(6))

                    except:
                        document = generar_titulo(document,'Tabla ' + pregunta['Tipo de escala'] + ' no especificada!')

                else: 
                    document = generar_titulo(document,'\n \n \n')



    delete_element_inseide_folder('docs/word/*')

    word_path = 'docs/word/'+ document_name + '.docx'
    document.save(word_path)

    #file_path = 'docs/demo.docx'
    #file_name = 'demo.docx'

    file_name = metadatos.cliente + '-' + datetime.today().strftime('%Y-%m-%d') + '.docx'

    documento_word = FileResponse(path=word_path, filename=file_name, media_type='application/octet-stream')

    return documento_word


def generar_titulo(document, titulo):

    table = document.add_table(rows=1, cols=1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = titulo
    document.add_paragraph(' ')
    return document




def get_unique_values_from_dict_list(list_of_dicts, value):

    result = []

    for i in range(len(list_of_dicts)):

        result.append(list_of_dicts[i][value])

    result = list(set(result))
    result.sort()

    return result 

def obtener_inidices_orden_de_las_preguntas(list_of_dicts):

    grupos = get_unique_values_from_dict_list(list_of_dicts, 'Grupos')

    result = {}

    for grupo in grupos:

        for value in list_of_dicts:

            if value['Grupos'] == grupo:
                
                result[grupo] = value['Categorias']
                break

    return list(result.values())

                



def int_to_Roman(num):
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
        ]
    syb = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
        ]
    roman_num = ''
    i = 0
    while  num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num

def delete_element_inseide_folder(folder_path):
    files = glob.glob(folder_path)
    for f in files:
        os.remove(f)


def normalizar_texto(txt):
    try:
        x = txt.lower().lstrip().rstrip()
        return x
    except:
        return txt

def id_generator(size=6, chars=string.ascii_uppercase + string.digits):

    return(''.join(random.choice(chars) for _ in range(size)))